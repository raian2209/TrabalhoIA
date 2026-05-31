"""Gera o artigo em .docx no estilo do template da SBC, preenchido com os
resultados reais do treino (outputs/metrics.json) e as figuras (figures/).

Uso:
    python gerar_artigo.py

Saida:
    artigo_sbc_mnist.docx
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
OUTPUTS = ROOT / "outputs"
FIGURES = ROOT / "figures"

FONT = "Times New Roman"


def load_metrics() -> dict:
    return json.loads((OUTPUTS / "metrics.json").read_text())


# --------------------------------------------------------------------------- #
# Helpers de formatacao no padrao SBC
# --------------------------------------------------------------------------- #
def set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = FONT
    p.paragraph_format.space_after = Pt(12)


def add_authors(doc: Document, authors: str, affiliation_lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(authors)
    r.font.size = Pt(12)
    r.font.name = FONT
    p.paragraph_format.space_after = Pt(6)
    for line in affiliation_lines:
        pa = doc.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ra = pa.add_run(line)
        ra.font.size = Pt(12)
        ra.font.name = FONT
        ra.italic = True
        pa.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_abstract(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rl = p.add_run(f"{label}. ")
    rl.bold = True
    rl.italic = True
    rl.font.name = FONT
    rl.font.size = Pt(12)
    rt = p.add_run(text)
    rt.italic = True
    rt.font.name = FONT
    rt.font.size = Pt(12)


def add_heading(doc: Document, number: str, title: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{number}. {title}" if number else title)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(13 if level == 1 else 12)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.5)
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(12)


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 12.0) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = FONT
    r.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(10)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    for line in code.strip("\n").split("\n"):
        r = p.add_run(line + "\n")
        r.font.name = "Courier New"
        r.font.size = Pt(9)


def add_bullet(doc: Document, text_bold: str, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text_bold)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(12)
    r2 = p.add_run(text)
    r2.font.name = FONT
    r2.font.size = Pt(12)


# --------------------------------------------------------------------------- #
# Construcao do documento
# --------------------------------------------------------------------------- #
def build(metrics: dict) -> Document:
    cfg = metrics["config"]
    acc = metrics["acuracia_teste_final"] * 100
    err = (1 - metrics["acuracia_teste_final"]) * 100
    params = cfg["parametros_treinaveis"]
    tempo = cfg["tempo_treino_s"]
    n_train = cfg["n_treino"]
    n_val = cfg["n_validacao"]
    n_test = cfg["n_teste"]
    val_acc = metrics["melhor_acuracia_validacao"] * 100

    def milhar(n: int) -> str:
        return f"{n:,}".replace(",", ".")
    rel = metrics["relatorio_por_classe"]
    # classe com menor F1 (mais dificil)
    pior = min(rel.items(), key=lambda kv: kv[1]["f1"])
    melhor = max(rel.items(), key=lambda kv: kv[1]["f1"])

    doc = Document()
    set_base_style(doc)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(3.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # --- Titulo e autoria ---
    add_title(doc, "Reconhecimento de Digitos Manuscritos com Redes Neurais "
                   "Convolucionais: um Estudo de Caso com a Base MNIST")
    add_authors(
        doc,
        "Raian Andrade",
        [
            "Curso de Ciencia da Computacao",
            "Disciplina de Inteligencia Artificial",
            "andradeneto1999@gmail.com",
        ],
    )

    add_abstract(
        doc,
        "Abstract",
        "This paper presents a case study on handwritten digit recognition using a "
        "convolutional neural network (CNN) trained on the MNIST dataset. We describe "
        "the network architecture, the supervised learning algorithm (backpropagation "
        "with the Adam optimizer and cross-entropy loss), the training data and the "
        f"experimental protocol. The model reaches {acc:.2f}% accuracy on the 10,000 "
        "test images, and we analyze its behavior through learning curves, a confusion "
        "matrix and per-class metrics. The full implementation in Python/PyTorch is "
        "discussed throughout the text.",
    )
    add_abstract(
        doc,
        "Resumo",
        "Este artigo apresenta um estudo de caso sobre reconhecimento de digitos "
        "manuscritos utilizando uma rede neural convolucional (CNN) treinada na base "
        "MNIST. Descrevemos a arquitetura da rede, o algoritmo de aprendizado "
        "supervisionado (retropropagacao com o otimizador Adam e funcao de perda de "
        "entropia cruzada), os dados de treino e o protocolo experimental. O modelo "
        f"atinge {acc:.2f}% de acuracia sobre as 10.000 imagens de teste, e analisamos "
        "seu comportamento por meio de curvas de aprendizado, matriz de confusao e "
        "metricas por classe. A implementacao completa em Python/PyTorch e discutida "
        "ao longo do texto.",
    )

    # --- 1. Introducao ---
    add_heading(doc, "1", "Introducao")
    add_body(
        doc,
        "O reconhecimento de digitos manuscritos e um problema classico de visao "
        "computacional e aprendizado de maquina, com aplicacoes diretas na leitura "
        "automatica de codigos postais, no processamento de cheques bancarios, na "
        "digitalizacao de formularios e na conversao de documentos manuscritos para "
        "formato digital. Apesar de parecer trivial para um ser humano, o problema "
        "envolve grande variabilidade na forma como cada pessoa escreve um mesmo "
        "digito, o que torna regras fixas inviaveis e motiva o uso de tecnicas de "
        "aprendizado a partir de exemplos.",
    )
    add_body(
        doc,
        "Neste trabalho conduzimos um estudo de caso sobre o tema utilizando a base "
        "MNIST como dominio de aplicacao e uma rede neural convolucional (CNN) como "
        "modelo de classificacao. O objetivo e demonstrar, de ponta a ponta, como uma "
        "rede neural aprende a reconhecer os dez digitos (0 a 9): desde o carregamento "
        "e pre-processamento dos dados, passando pela definicao da arquitetura e do "
        "algoritmo de treino, ate a avaliacao quantitativa dos resultados. Toda a "
        "discussao e acompanhada da implementacao em Python com a biblioteca PyTorch.",
    )
    add_body(
        doc,
        "O restante do artigo esta organizado da seguinte forma. A Secao 2 descreve o "
        "dominio de aplicacao e a base de dados MNIST. A Secao 3 apresenta a "
        "arquitetura da rede neural. A Secao 4 detalha o algoritmo de aprendizado. A "
        "Secao 5 descreve a implementacao em codigo. A Secao 6 apresenta e discute os "
        "resultados experimentais, e a Secao 7 conclui o trabalho.",
    )

    # --- 2. Dominio e base de dados ---
    add_heading(doc, "2", "Dominio de Aplicacao e Base de Dados")
    add_body(
        doc,
        "O dominio de aplicacao deste estudo e a classificacao de imagens de digitos "
        "manuscritos isolados. Utilizamos a base MNIST (Modified National Institute of "
        "Standards and Technology), um dos conjuntos de dados mais utilizados como "
        "referencia em aprendizado de maquina. A MNIST e composta por 70.000 imagens "
        "em escala de cinza, sendo 60.000 destinadas ao treino e 10.000 ao teste. Cada "
        "imagem possui 28x28 pixels (784 atributos de intensidade no intervalo de 0 a "
        "255) e esta associada a um rotulo inteiro entre 0 e 9, indicando o digito "
        "correspondente.",
    )
    add_body(
        doc,
        "As imagens ja vem centralizadas e normalizadas em tamanho, o que reduz parte "
        "da variabilidade e permite concentrar o estudo no aprendizado das formas dos "
        "digitos. Como pre-processamento adicional, convertemos cada imagem para um "
        "tensor e aplicamos a normalizacao padrao da base, subtraindo a media (0,1307) "
        "e dividindo pelo desvio-padrao (0,3081) calculados sobre o conjunto de treino. "
        "Essa padronizacao centraliza os dados em torno de zero e acelera a "
        "convergencia do treino. A Figura 1 ilustra algumas amostras da base.",
    )
    add_figure(
        doc,
        FIGURES / "amostras.png",
        "Figura 1. Amostras de digitos manuscritos da base MNIST (28x28 pixels, "
        "escala de cinza) com seus respectivos rotulos.",
        width_cm=12.5,
    )
    add_heading(doc, "2.1", "Divisao dos Dados (Treino, Validacao e Teste)", level=2)
    add_body(
        doc,
        "Para uma avaliacao metodologicamente correta, adotamos uma particao em tres "
        f"conjuntos. As 60.000 imagens originais de treino foram divididas em "
        f"{milhar(n_train)} imagens de treino e {milhar(n_val)} imagens de validacao, "
        "por meio de uma particao aleatoria reproduzivel (semente fixa). As "
        f"{milhar(n_test)} imagens do conjunto oficial de teste foram mantidas "
        "intactas. O conjunto de treino e usado para ajustar os pesos da rede; o "
        "conjunto de validacao serve para monitorar a generalizacao a cada epoca e "
        "selecionar o melhor modelo (early model selection), evitando que decisoes de "
        "projeto sejam influenciadas pelo teste; e o conjunto de teste e utilizado "
        "uma unica vez, ao final, para estimar de forma imparcial o desempenho do "
        "modelo escolhido. A Tabela 1 resume essa divisao.",
    )
    add_split_table(doc, n_train, n_val, n_test)

    # --- 3. Arquitetura ---
    add_heading(doc, "3", "Arquitetura da Rede Neural")
    add_body(
        doc,
        "Para a tarefa de classificacao adotamos uma rede neural convolucional (CNN), "
        "arquitetura especialmente adequada a dados com estrutura espacial como "
        "imagens. Diferentemente de uma rede totalmente conectada, a CNN explora a "
        "correlacao local entre pixels vizinhos por meio de filtros convolucionais "
        "compartilhados, o que reduz drasticamente o numero de parametros e introduz "
        "invariancia a pequenas translacoes dos tracos do digito.",
    )
    add_body(
        doc,
        "A arquitetura proposta, denominada DigitCNN, e composta por dois blocos "
        "convolucionais seguidos de um classificador totalmente conectado. Cada bloco "
        "aplica duas camadas convolucionais com filtros 3x3, normalizacao em lote "
        "(BatchNorm), ativacao ReLU e uma reducao espacial por max-pooling 2x2. O "
        "primeiro bloco eleva a profundidade de 1 para 32 canais e reduz a resolucao "
        "de 28x28 para 14x14; o segundo eleva de 32 para 64 canais e reduz de 14x14 "
        "para 7x7. Em seguida, o mapa de caracteristicas resultante (64x7x7) e achatado "
        "e processado por uma camada densa de 128 neuronios com Dropout de 0,3, "
        "finalizando em uma camada de saida com 10 neuronios (um por digito). A Tabela "
        "2 resume a arquitetura.",
    )
    add_arch_table(doc, params)
    add_body(
        doc,
        "A camada de saida produz dez valores (logits), convertidos em uma "
        "distribuicao de probabilidade pela funcao softmax durante a inferencia. O "
        "digito predito e aquele de maior probabilidade. No total, a rede possui "
        f"{params:,} parametros treinaveis, o que a torna leve o suficiente para ser "
        "treinada em poucos minutos em CPU.".replace(",", "."),
    )

    # --- 4. Algoritmo de aprendizado ---
    add_heading(doc, "4", "Algoritmo de Aprendizado")
    add_body(
        doc,
        "O treinamento segue o paradigma de aprendizado supervisionado: a rede recebe "
        "pares (imagem, rotulo) e ajusta seus parametros para minimizar o erro entre a "
        "saida prevista e o rotulo verdadeiro. A funcao de perda utilizada e a entropia "
        "cruzada (cross-entropy), apropriada para classificacao multiclasse, que mede a "
        "discrepancia entre a distribuicao de probabilidade prevista e o rotulo real "
        "codificado.",
    )
    add_body(
        doc,
        "O ajuste dos pesos e feito pelo algoritmo de retropropagacao (backpropagation), "
        "que calcula o gradiente da perda em relacao a cada parametro por meio da regra "
        "da cadeia, propagando o erro da saida ate as primeiras camadas. A atualizacao "
        "dos pesos emprega o otimizador Adam, uma variante do gradiente descendente "
        "estocastico que adapta a taxa de aprendizado para cada parametro com base em "
        "estimativas dos momentos de primeira e segunda ordem do gradiente, "
        "proporcionando convergencia rapida e estavel.",
    )
    add_body(doc, "Os principais hiperparametros do treino foram:")
    add_bullet(doc, "Otimizador: ", f"{cfg['otimizador']} com taxa de aprendizado "
               f"(learning rate) de {cfg['lr']}.")
    add_bullet(doc, "Funcao de perda: ", "entropia cruzada (CrossEntropyLoss).")
    add_bullet(doc, "Tamanho do lote (batch size): ", f"{cfg['batch_size']} imagens.")
    add_bullet(doc, "Numero de epocas: ", f"{cfg['epocas']} passagens completas pelo "
               "conjunto de treino.")
    add_bullet(doc, "Semente aleatoria: ", f"{cfg['seed']} (para reprodutibilidade).")
    add_body(
        doc,
        "A cada epoca, o conjunto de treino e percorrido em lotes embaralhados; para "
        "cada lote calcula-se a saida da rede (forward), a perda, os gradientes "
        "(backward) e a atualizacao dos pesos. Ao final de cada epoca, o modelo e "
        "avaliado no conjunto de validacao, e o checkpoint de maior acuracia de "
        "validacao e preservado. O conjunto de teste nao participa em nenhum momento "
        "do treino ou da selecao de modelo, sendo reservado para a avaliacao final.",
    )

    # --- 5. Implementacao ---
    add_heading(doc, "5", "Implementacao")
    add_body(
        doc,
        "A implementacao foi feita em Python utilizando a biblioteca PyTorch para a "
        "definicao e o treino da rede, e torchvision para o acesso a base MNIST. O "
        "codigo esta organizado em modulos: data.py (carregamento e pre-processamento), "
        "model.py (arquitetura), train.py (laco de treino e avaliacao) e visualize.py "
        "(geracao das figuras). O trecho a seguir mostra a definicao da arquitetura.",
    )
    add_code(
        doc,
        """class DigitCNN(nn.Module):
    def __init__(self, num_classes=10):
        super().__init__()
        self.features = nn.Sequential(
            nn.Conv2d(1, 32, 3, padding=1), nn.BatchNorm2d(32), nn.ReLU(),
            nn.Conv2d(32, 32, 3, padding=1), nn.BatchNorm2d(32), nn.ReLU(),
            nn.MaxPool2d(2),                              # 28x28 -> 14x14
            nn.Conv2d(32, 64, 3, padding=1), nn.BatchNorm2d(64), nn.ReLU(),
            nn.Conv2d(64, 64, 3, padding=1), nn.BatchNorm2d(64), nn.ReLU(),
            nn.MaxPool2d(2))                              # 14x14 -> 7x7
        self.classifier = nn.Sequential(
            nn.Flatten(), nn.Linear(64*7*7, 128), nn.ReLU(),
            nn.Dropout(0.3), nn.Linear(128, num_classes))

    def forward(self, x):
        return self.classifier(self.features(x))""",
    )
    add_body(doc, "O laco de treino de uma epoca resume-se a:")
    add_code(
        doc,
        """for x, y in train_loader:
    optimizer.zero_grad()
    logits = model(x)               # forward
    loss = criterion(logits, y)     # entropia cruzada
    loss.backward()                 # backpropagation
    optimizer.step()                # atualizacao (Adam)""",
    )

    # --- 6. Resultados ---
    add_heading(doc, "6", "Resultados e Discussao")
    add_body(
        doc,
        f"Apos {cfg['epocas']} epocas de treino (concluidas em aproximadamente "
        f"{tempo:.0f} segundos em CPU), o modelo selecionado pela validacao (melhor "
        f"acuracia de validacao: {val_acc:.2f}%) alcancou {acc:.2f}% de acuracia sobre "
        f"as {milhar(n_test)} imagens do conjunto de teste, o que corresponde a uma "
        f"taxa de erro de apenas {err:.2f}%. A Figura 2 apresenta a evolucao da perda e "
        "da acuracia ao longo das epocas, para os conjuntos de treino e validacao. "
        "Observa-se convergencia rapida e a ausencia de overfitting acentuado, com as "
        "curvas de treino e validacao proximas entre si.",
    )
    add_figure(
        doc,
        FIGURES / "curvas.png",
        "Figura 2. Curvas de perda (esquerda) e acuracia (direita) por epoca, para os "
        "conjuntos de treino e validacao.",
        width_cm=14.0,
    )
    add_body(
        doc,
        "Para uma analise mais detalhada dos acertos e erros por classe, a Figura 3 "
        "exibe a matriz de confusao sobre o conjunto de teste. A forte concentracao de "
        "valores na diagonal principal confirma o bom desempenho do classificador. Os "
        "poucos erros ocorrem majoritariamente entre digitos de formato visualmente "
        "semelhante.",
    )
    add_figure(
        doc,
        FIGURES / "matriz_confusao.png",
        "Figura 3. Matriz de confusao sobre as 10.000 imagens de teste. Linhas "
        "representam a classe verdadeira e colunas a classe predita.",
        width_cm=11.0,
    )
    add_body(
        doc,
        "A Tabela 3 apresenta as metricas de precisao, revocacao e F1-score por classe. "
        f"O digito '{melhor[0]}' obteve o melhor F1-score ({melhor[1]['f1']*100:.1f}%), "
        f"enquanto o digito '{pior[0]}' foi o mais dificil ({pior[1]['f1']*100:.1f}% de "
        "F1-score), coerente com sua maior semelhanca grafica com outros digitos.",
    )
    add_metrics_table(doc, rel)
    add_body(
        doc,
        "Por fim, a Figura 4 mostra exemplos de predicoes da rede, com a probabilidade "
        "atribuida a classe escolhida. A grande maioria das predicoes e correta e com "
        "alta confianca, ilustrando qualitativamente a eficacia do modelo.",
    )
    add_figure(
        doc,
        FIGURES / "predicoes.png",
        "Figura 4. Exemplos de predicoes da CNN. Em verde os acertos e em vermelho os "
        "eventuais erros, com a probabilidade da classe predita.",
        width_cm=14.0,
    )

    # --- 7. Conclusao ---
    add_heading(doc, "7", "Conclusao")
    add_body(
        doc,
        "Este estudo de caso demonstrou, de forma completa, a aplicacao de uma rede "
        "neural convolucional ao reconhecimento de digitos manuscritos da base MNIST. "
        "A partir de uma arquitetura compacta e de um treino supervisionado com "
        "backpropagation e otimizador Adam, o modelo atingiu uma acuracia de "
        f"{acc:.2f}% no conjunto de teste, comprovando a eficacia das CNNs para tarefas "
        "de classificacao de imagens mesmo com recursos computacionais modestos.",
    )
    add_body(
        doc,
        "Como trabalhos futuros, destacam-se a aplicacao de tecnicas de aumento de "
        "dados (data augmentation), a avaliacao da robustez do modelo frente a "
        "perturbacoes adversariais e a extensao da abordagem para bases mais "
        "desafiadoras, como EMNIST e dados de digitos do mundo real.",
    )

    # --- Referencias ---
    add_heading(doc, "", "Referencias")
    for ref in [
        "LeCun, Y., Bottou, L., Bengio, Y. and Haffner, P. (1998). Gradient-based "
        "learning applied to document recognition. Proceedings of the IEEE, "
        "86(11):2278-2324.",
        "Goodfellow, I., Bengio, Y. and Courville, A. (2016). Deep Learning. MIT Press.",
        "Kingma, D. P. and Ba, J. (2015). Adam: A Method for Stochastic Optimization. "
        "International Conference on Learning Representations (ICLR).",
        "Paszke, A. et al. (2019). PyTorch: An Imperative Style, High-Performance Deep "
        "Learning Library. Advances in Neural Information Processing Systems (NeurIPS).",
    ]:
        p = doc.add_paragraph(ref)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.0)
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.name = FONT
            r.font.size = Pt(12)

    return doc


def add_split_table(doc: Document, n_train: int, n_val: int, n_test: int) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run("Tabela 1. Divisao da base MNIST em treino, validacao e teste.")
    rc.font.name = FONT
    rc.font.size = Pt(10)
    total = n_train + n_val + n_test

    def fmt(n: int) -> str:
        return f"{n:,}".replace(",", ".")

    rows = [
        ("Conjunto", "Imagens", "Proporcao", "Uso"),
        ("Treino", fmt(n_train), f"{100*n_train/total:.0f}%", "ajuste dos pesos"),
        ("Validacao", fmt(n_val), f"{100*n_val/total:.0f}%", "selecao do modelo"),
        ("Teste", fmt(n_test), f"{100*n_test/total:.0f}%", "avaliacao final"),
        ("Total", fmt(total), "100%", "-"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.cell(i, j)
            cell.text = txt
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.name = FONT
                    r.font.size = Pt(10)
                    if i == 0 or (i == len(rows) - 1):
                        r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_arch_table(doc: Document, params: int) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run("Tabela 2. Arquitetura da rede DigitCNN.")
    rc.font.name = FONT
    rc.font.size = Pt(10)
    rows = [
        ("Camada", "Configuracao", "Saida"),
        ("Entrada", "imagem em escala de cinza", "1 x 28 x 28"),
        ("Conv + BN + ReLU (x2)", "32 filtros 3x3", "32 x 28 x 28"),
        ("MaxPool 2x2", "reducao espacial", "32 x 14 x 14"),
        ("Conv + BN + ReLU (x2)", "64 filtros 3x3", "64 x 14 x 14"),
        ("MaxPool 2x2", "reducao espacial", "64 x 7 x 7"),
        ("Flatten + Linear + ReLU", "128 neuronios", "128"),
        ("Dropout", "p = 0,3", "128"),
        ("Linear (saida)", "10 neuronios (digitos)", "10"),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.cell(i, j)
            cell.text = txt
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.name = FONT
                    r.font.size = Pt(10)
                    if i == 0:
                        r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_metrics_table(doc: Document, rel: dict) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run("Tabela 3. Metricas por classe no conjunto de teste (%).")
    rc.font.name = FONT
    rc.font.size = Pt(10)
    headers = ("Digito", "Precisao", "Revocacao", "F1-score", "Suporte")
    table = doc.add_table(rows=len(rel) + 1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = h
        for par in c.paragraphs:
            for r in par.runs:
                r.font.name = FONT
                r.font.size = Pt(10)
                r.bold = True
    for i, (digit, m) in enumerate(rel.items(), start=1):
        vals = [
            digit,
            f"{m['precisao']*100:.1f}",
            f"{m['revocacao']*100:.1f}",
            f"{m['f1']*100:.1f}",
            str(m["suporte"]),
        ]
        for j, v in enumerate(vals):
            c = table.cell(i, j)
            c.text = v
            for par in c.paragraphs:
                for r in par.runs:
                    r.font.name = FONT
                    r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def main() -> None:
    metrics = load_metrics()
    doc = build(metrics)
    out = ROOT / "artigo_sbc_mnist.docx"
    doc.save(str(out))
    print(f"Artigo gerado em {out}")


if __name__ == "__main__":
    main()
